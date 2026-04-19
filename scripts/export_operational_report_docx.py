from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = PROJECT_ROOT / "docs" / "operational_database_report.md"
TARGET_PATH = PROJECT_ROOT / "docs" / "operational_database_report.docx"


def add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return

    if stripped.startswith("### "):
        document.add_heading(stripped[4:], level=3)
        return
    if stripped.startswith("## "):
        document.add_heading(stripped[3:], level=2)
        return
    if stripped.startswith("# "):
        document.add_heading(stripped[2:], level=1)
        return
    if stripped.startswith("- "):
        document.add_paragraph(stripped[2:], style="List Bullet")
        return

    document.add_paragraph(stripped)


def main() -> None:
    document = Document()
    for line in SOURCE_PATH.read_text(encoding="utf-8").splitlines():
        add_markdown_line(document, line)
    document.save(TARGET_PATH)
    print(f"Word report exported to: {TARGET_PATH}")


if __name__ == "__main__":
    main()
